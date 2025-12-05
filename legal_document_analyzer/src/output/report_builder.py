
from pathlib import Path
import json
from typing import Dict, List

from ..utils.file_utils import write_text


def to_json(outdir: str | Path, data: Dict) -> Path:
    p = Path(outdir) / 'analysis.json'
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(json.dumps(data, indent=2), encoding='utf-8')
    return p


def to_markdown(outdir: str | Path, summary: str, clauses: List[Dict], risk: Dict) -> Path:
    md = ['# Legal Document Analysis Report', '
']
    md += ['## Summary', '', summary, '
']
    md += ['## Extracted Clauses', '']
    for cl in clauses:
        md += [f"### {cl['type'].replace('_',' ').title()} ({cl['section']})", '', cl['text'], '']
    md += ['## Risk Analysis', '', f"Score: {risk['score']} | Level: **{risk['level'].upper()}**", '']
    for f in risk.get('findings', []):
        md += [f"- {f['clause_type'].title()} → " + ', '.join([r['rule'] for r in f['risks']])]
    return write_text(Path(outdir) / 'analysis.md', '
'.join(md))


def to_docx(outdir: str | Path, summary: str, clauses: List[Dict], risk: Dict) -> Path:
    try:
        import docx
        doc = docx.Document()
        doc.add_heading('Legal Document Analysis Report', level=1)
        doc.add_heading('Summary', level=2)
        doc.add_paragraph(summary)
        doc.add_heading('Extracted Clauses', level=2)
        for cl in clauses:
            doc.add_heading(f"{cl['type'].replace('_',' ').title()} ({cl['section']})", level=3)
            doc.add_paragraph(cl['text'])
        doc.add_heading('Risk Analysis', level=2)
        doc.add_paragraph(f"Score: {risk['score']} | Level: {risk['level'].upper()}")
        for f in risk.get('findings', []):
            doc.add_paragraph(f"{f['clause_type'].title()} → " + ', '.join([r['rule'] for r in f['risks']]))
        out = Path(outdir) / 'analysis.docx'
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(out)
        return out
    except Exception as e:
        # Fallback to markdown
        return to_markdown(outdir, summary, clauses, risk)
